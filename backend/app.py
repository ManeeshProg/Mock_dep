from fastapi import FastAPI, UploadFile, File, Form
from fastapi import HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn
from typing import List, Optional
import io
from fastapi import WebSocket, WebSocketDisconnect
import tempfile
import os
from sarvamai import SarvamAI

from fastapi.responses import StreamingResponse

# PDF / DOCX generation
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
import docx
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = FastAPI(title="Resume Savvy RAG API")

# Enable CORS for development (allow frontend dev server to call API)
app.add_middleware(
	CORSMiddleware,
	allow_origins=["*"],
	allow_credentials=True,
	allow_methods=["*"],
	allow_headers=["*"],
)


# Restore correct QuestionsRequest and EvaluateRequest definitions

class QuestionsRequest(BaseModel):
	session_id: str
	role: Optional[str] = None
	count_role: int = 7
	count_resume: int = 8

class EvaluateRequest(BaseModel):
	session_id: str
	role: str
	technical_answers: List[dict]
	hr_answers: List[dict]



class ReportRequest(BaseModel):
	session_id: str
	candidate_name: Optional[str] = None
	role: Optional[str] = None
	# Optional detailed answers (legacy) - not used when strengths/improvements provided
	technical_answers: List[dict] = []
	hr_answers: List[dict] = []
	# Concise performance fields
	overall: Optional[float] = None
	technical_score: Optional[float] = None  # percentage (optional)
	hr_score: Optional[float] = None  # percentage (optional)
	role_score: Optional[float] = None  # percentage (optional)
	resume_score: Optional[float] = None  # percentage (optional)
	# Actual marks scored (use these if provided)
	role_marks: Optional[float] = None  # marks scored out of 35
	resume_marks: Optional[float] = None  # marks scored out of 40
	technical_marks: Optional[float] = None  # marks scored out of 75 (role+resume)
	hr_marks: Optional[float] = None  # marks scored out of 25
	total_marks: Optional[float] = None  # marks scored out of 100
	strengths: List[str] = []
	improvements: List[str] = []

	# LLM-generated detailed feedback sections (preferred over strengths/improvements)
	technical_feedback: List[str] = []
	hr_feedback: List[str] = []
	communication_feedback: List[str] = []
	tips_to_improve: List[str] = []



# Restore ExtractResponse definition and rag_engine import
class ExtractResponse(BaseModel):
	session_id: str
	chunks_indexed: int
	metadata: dict

from rag.engine import RAGEngine
rag_engine = RAGEngine()

@app.post("/extract", response_model=ExtractResponse)
async def extract_resume(session_id: str = Form(...), file: UploadFile = File(...)):
	content_text, metadata = await rag_engine.extract_and_index(session_id=session_id, upload_file=file)
	return ExtractResponse(session_id=session_id, chunks_indexed=metadata.get("chunks_indexed", 0), metadata=metadata)



@app.post("/questions/technical")
async def generate_technical_questions(body: QuestionsRequest):
	try:
		questions = await rag_engine.generate_technical_questions(
			session_id=body.session_id,
			role=body.role or "Full Stack Developer",
			count_role=body.count_role,
			count_resume=body.count_resume,
		)
		return {"questions": questions}
	except Exception as e:
		raise HTTPException(status_code=503, detail=f"Question generation unavailable: {e}")



@app.post("/questions/hr")
async def generate_hr_questions(body: QuestionsRequest):
	try:
		questions = await rag_engine.generate_hr_questions(
			session_id=body.session_id,
			count=5,
		)
		return {"questions": questions}
	except Exception as e:
		raise HTTPException(status_code=503, detail=f"Question generation unavailable: {e}")



@app.post("/evaluate")
async def evaluate(body: EvaluateRequest):
	try:
		result = await rag_engine.evaluate_answers(
			session_id=body.session_id,
			role=body.role,
			technical_answers=body.technical_answers,
			hr_answers=body.hr_answers,
		)
		return result
	except Exception as e:
		raise HTTPException(status_code=503, detail=f"Evaluation unavailable: {e}")


@app.post("/report")
async def generate_report(body: ReportRequest, format: Optional[str] = "pdf"):
	"""Generate and return interview report as PDF or DOCX.

	The request body should contain interview data (session id, role, answers, and optional scores).
	Query param `format` accepts `pdf` or `docx` (default `pdf`).
	"""
	fmt = (format or "pdf").lower()

	# Build textual content for the report (concise performance summary preferred)
	candidate_name = body.candidate_name or "Candidate"
	title = f"{candidate_name} - Interview Results"
	subtitle = f"Role: {body.role or 'N/A'}"
	overall = body.overall if body.overall is not None else round(((body.technical_score or 0) + (body.hr_score or 0)) / 2, 2)

	strengths = body.strengths or []
	improvements = body.improvements or []
	
	# Generate HR performance description
	hr_performance = ""
	if body.hr_score is not None:
		hr_score = body.hr_score
		if hr_score >= 85:
			hr_performance = "Excellent communication and interpersonal skills demonstrated throughout the interview. The candidate articulated responses clearly, showed strong cultural fit, and demonstrated strong leadership potential and emotional intelligence."
		elif hr_score >= 75:
			hr_performance = "Good communication and behavioral responses. The candidate showed reasonable clarity in expressing ideas, demonstrated team collaboration mindset, and positive attitude towards learning and growth."
		elif hr_score >= 65:
			hr_performance = "Moderate communication skills with some room for improvement. The candidate communicated key points but could enhance clarity and storytelling. STAR method practice would help in articulating experiences more effectively."
		elif hr_score >= 50:
			hr_performance = "Communication and behavioral skills need significant improvement. The candidate should focus on developing clearer articulation of thoughts, using structured storytelling (STAR method), and demonstrating stronger emotional intelligence in interactions."
		else:
			hr_performance = "Significant improvement needed in communication, interpersonal, and behavioral areas. Consider working with a communication coach, practicing mock interviews, and focusing on developing soft skills and confidence."
	
	# Generate "Tips to Enhance Knowledge" from low-scoring areas AND answer feedback
	tips = []
	
	# Analyze technical performance
	if body.technical_score is not None and body.technical_score < 60:
		tips.append("Focus on mastering fundamental data structures (arrays, linked lists, trees, hash maps) and their operations.")
	elif body.technical_score is not None and body.technical_score < 75:
		tips.append("Practice coding problems on LeetCode or HackerRank focusing on your weak areas (arrays, strings, sorting).")
	
	# Role-specific recommendations
	if body.role:
		role_lower = body.role.lower()
		if "python" in role_lower:
			tips.append("Strengthen Python fundamentals: list comprehensions, decorators, generators, and async/await patterns.")
		elif "javascript" in role_lower or "react" in role_lower:
			tips.append("Deepen knowledge of JavaScript closures, async/await, promises, and React hooks lifecycle.")
		elif "java" in role_lower:
			tips.append("Master Java concepts: generics, exception handling, multithreading, and Spring framework basics.")
		elif "full stack" in role_lower or "developer" in role_lower:
			tips.append("Study both frontend (HTML/CSS/JS) and backend fundamentals (APIs, databases, authentication).")
		else:
			tips.append(f"Study core concepts and best practices specific to {body.role} role.")
	
	# HR/Behavioral performance
	if body.hr_score is not None and body.hr_score < 60:
		tips.append("Work on communication and presentation skills - practice articulating your thoughts clearly and concisely.")
	elif body.hr_score is not None and body.hr_score < 75:
		tips.append("Improve storytelling: use STAR method (Situation, Task, Action, Result) to explain past experiences effectively.")
	
	# General improvement tips
	if not tips or len(tips) < 2:
		if body.technical_score is not None and body.technical_score > 0:
			tips.append("Continue practicing coding problems and system design concepts to strengthen technical foundation.")
		tips.append("Stay updated with industry trends, read technical blogs, and contribute to open source projects.")

	try:
		if fmt == "pdf":
			from reportlab.platypus import Table, TableStyle
			from reportlab.lib import colors
			from reportlab.lib.units import inch
			
			buf = io.BytesIO()
			doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
			styles = getSampleStyleSheet()
			story = []

			story.append(Paragraph(title, styles["Title"]))
			story.append(Paragraph(subtitle, styles["Normal"]))
			story.append(Spacer(1, 12))

			# Score Table
			tech_score = body.technical_score if body.technical_score is not None else 0
			hr_score = body.hr_score if body.hr_score is not None else 0
			role_score = body.role_score if getattr(body, 'role_score', None) is not None else 0
			resume_score = body.resume_score if getattr(body, 'resume_score', None) is not None else 0
			
			# Get marks (calculate if not provided) - mapping: role 7->35, resume 8->40, hr 5->25
			role_marks = body.role_marks if body.role_marks is not None else (role_score * 35 / 100)
			resume_marks = body.resume_marks if body.resume_marks is not None else (resume_score * 40 / 100)
			technical_marks = body.technical_marks if body.technical_marks is not None else (role_marks + resume_marks)
			hr_marks = body.hr_marks if body.hr_marks is not None else (hr_score * 25 / 100)
			total_marks = body.total_marks if body.total_marks is not None else (technical_marks + hr_marks)
			
			# Create score table data: Category | Score (marks) | Total | Percentage
			table_data = [
				["Category", "Score", "Total Score", "Percentage"],
			]
			if role_score:
				table_data.append(["Role-based Technical", f"{int(role_marks)}", "35", f"{int(role_score)}%"])
			if resume_score:
				table_data.append(["Resume-based Technical", f"{int(resume_marks)}", "40", f"{int(resume_score)}%"])
			if tech_score:
				table_data.append(["Overall Technical", f"{int(technical_marks)}", "75", f"{int(tech_score)}%"])
			if hr_score:
				table_data.append(["HR & Behavioral", f"{int(hr_marks)}", "25", f"{int(hr_score)}%"])
			table_data.append(["Total Score", f"{int(total_marks)}", "100", f"{int(overall)}%"])
			
			table = Table(table_data, colWidths=[2*inch, 1*inch, 1.2*inch, 1.2*inch])
			table.setStyle(TableStyle([
				('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D3AF37')),
				('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
				('ALIGN', (0, 0), (-1, -1), 'CENTER'),
				('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
				('FONTSIZE', (0, 0), (-1, 0), 10),
				('BOTTOMPADDING', (0, 0), (-1, 0), 12),
				('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#F0E68C')),
				('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
				('GRID', (0, 0), (-1, -1), 1, colors.black),
				('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#F5F5F5')]),
			]))
			story.append(table)
			story.append(Spacer(1, 16))

			# Performance in HR Interview Section
			if hr_performance:
				story.append(Paragraph("Performance in HR Interview", styles["Heading2"]))
				story.append(Paragraph(hr_performance, styles["BodyText"]))
				story.append(Spacer(1, 12))

			# Strengths
			if strengths:
				story.append(Paragraph("Strengths", styles["Heading2"]))
				for i, s in enumerate(strengths[:3], 1):
					story.append(Paragraph(f"{i}. {s}", styles["BodyText"]))
				story.append(Spacer(1, 12))

			# If detailed LLM feedback sections are provided, prefer them over generic strengths/improvements
			tech_fb = body.technical_feedback or []
			hr_fb = body.hr_feedback or []
			comm_fb = body.communication_feedback or []
			tips_fb = body.tips_to_improve or []

			if tech_fb:
				story.append(Paragraph("Technical Feedback", styles["Heading2"]))
				for i, s in enumerate(tech_fb[:4], 1):
					story.append(Paragraph(f"{i}. {s}", styles["BodyText"]))
				story.append(Spacer(1, 12))

			if hr_fb:
				story.append(Paragraph("HR / Behavioral Feedback", styles["Heading2"]))
				for i, s in enumerate(hr_fb[:4], 1):
					story.append(Paragraph(f"{i}. {s}", styles["BodyText"]))
				story.append(Spacer(1, 12))

			if comm_fb:
				story.append(Paragraph("Communication Feedback", styles["Heading2"]))
				for i, s in enumerate(comm_fb[:4], 1):
					story.append(Paragraph(f"{i}. {s}", styles["BodyText"]))
				story.append(Spacer(1, 12))

			if tips_fb:
				story.append(Paragraph("Tips to Improve", styles["Heading2"]))
				for i, s in enumerate(tips_fb[:4], 1):
					story.append(Paragraph(f"{i}. {s}", styles["BodyText"]))
				story.append(Spacer(1, 12))

			# Areas to Improve
			if improvements:
				story.append(Paragraph("Areas for Improvement", styles["Heading2"]))
				for i, it in enumerate(improvements[:3], 1):
					story.append(Paragraph(f"{i}. {it}", styles["BodyText"]))
				story.append(Spacer(1, 12))

			# Tips to Enhance Knowledge
			if tips:
				story.append(Paragraph("Tips to Enhance Knowledge", styles["Heading2"]))
				for i, tip in enumerate(tips[:3], 1):
					story.append(Paragraph(f"{i}. {tip}", styles["BodyText"]))
				story.append(Spacer(1, 12))

			doc.build(story)
			buf.seek(0)
			return StreamingResponse(buf, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=interview-report-{body.session_id}.pdf"})

		elif fmt == "docx":
			# Build a nicely formatted Word document using python-docx
			doc = docx.Document()
			doc.add_heading(title, level=0)

			if subtitle:
				doc.add_paragraph(subtitle)
			
		# Add scores table
		table = doc.add_table(rows=1, cols=4)
		table.style = 'Light Grid Accent 1'
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = "Category"
		hdr_cells[1].text = "Score"
		hdr_cells[2].text = "Total Score"
		hdr_cells[3].text = "Percentage"
		
		tech_score = body.technical_score if body.technical_score is not None else 0
		hr_score = body.hr_score if body.hr_score is not None else 0
		role_score = body.role_score if getattr(body, 'role_score', None) is not None else 0
		resume_score = body.resume_score if getattr(body, 'resume_score', None) is not None else 0
		
		# Get marks (calculate if not provided) - mapping: role 7->35, resume 8->40, hr 5->25
		role_marks = body.role_marks if body.role_marks is not None else (role_score * 35 / 100)
		resume_marks = body.resume_marks if body.resume_marks is not None else (resume_score * 40 / 100)
		technical_marks = body.technical_marks if body.technical_marks is not None else (role_marks + resume_marks)
		hr_marks = body.hr_marks if body.hr_marks is not None else (hr_score * 25 / 100)
		total_marks = body.total_marks if body.total_marks is not None else (technical_marks + hr_marks)
		
		if role_score:
			row_cells = table.add_row().cells
			row_cells[0].text = "Role-based Technical"
			row_cells[1].text = f"{int(role_marks)}"
			row_cells[2].text = "35"
			row_cells[3].text = f"{int(role_score)}%"
		if resume_score:
			row_cells = table.add_row().cells
			row_cells[0].text = "Resume-based Technical"
			row_cells[1].text = f"{int(resume_marks)}"
			row_cells[2].text = "40"
			row_cells[3].text = f"{int(resume_score)}%"
		if tech_score:
			row_cells = table.add_row().cells
			row_cells[0].text = "Overall Technical"
			row_cells[1].text = f"{int(technical_marks)}"
			row_cells[2].text = "75"
			row_cells[3].text = f"{int(tech_score)}%"
		if hr_score:
			row_cells = table.add_row().cells
			row_cells[0].text = "HR & Behavioral"
			row_cells[1].text = f"{int(hr_marks)}"
			row_cells[2].text = "25"
			row_cells[3].text = f"{int(hr_score)}%"
		
		row_cells = table.add_row().cells
		row_cells[0].text = "Total Score"
		row_cells[1].text = f"{int(total_marks)}"
		row_cells[2].text = "100"
		row_cells[3].text = f"{int(overall)}%"

		# Performance in HR Interview section
		if hr_performance:
			doc.add_paragraph()
			doc.add_heading("Performance in HR Interview", level=2)
			doc.add_paragraph(hr_performance)

			# If detailed LLM feedback sections are provided, prefer them
			tech_fb = body.technical_feedback or []
			hr_fb = body.hr_feedback or []
			comm_fb = body.communication_feedback or []
			tips_fb = body.tips_to_improve or []

			# Technical Feedback
			if tech_fb:
				doc.add_paragraph()
				doc.add_heading("Technical Feedback", level=2)
				for i, s in enumerate(tech_fb[:4], 1):
					doc.add_paragraph(f"{i}. {s}")

			# HR / Behavioral Feedback
			if hr_fb:
				doc.add_paragraph()
				doc.add_heading("HR / Behavioral Feedback", level=2)
				for i, s in enumerate(hr_fb[:4], 1):
					doc.add_paragraph(f"{i}. {s}")

			# Communication Feedback
			if comm_fb:
				doc.add_paragraph()
				doc.add_heading("Communication Feedback", level=2)
				for i, s in enumerate(comm_fb[:4], 1):
					doc.add_paragraph(f"{i}. {s}")

			# Tips to Improve
			if tips_fb:
				doc.add_paragraph()
				doc.add_heading("Tips to Improve", level=2)
				for i, s in enumerate(tips_fb[:4], 1):
					doc.add_paragraph(f"{i}. {s}")

			# Areas to Improve
			if improvements:
				doc.add_paragraph()
				doc.add_heading("Areas for Improvement", level=2)
				for i, it in enumerate(improvements[:3], 1):
					doc.add_paragraph(f"{i}. {it}")
			
			# Tips to Enhance Knowledge
			if tips:
				doc.add_paragraph()
				doc.add_heading("Tips to Enhance Knowledge", level=2)
				for i, tip in enumerate(tips[:3], 1):
					doc.add_paragraph(f"{i}. {tip}")

			# Finalize into BytesIO and return
			buf = io.BytesIO()
			doc.save(buf)
			buf.seek(0)
			return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=interview-report-{body.session_id}.docx"})

		else:
			raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'.")
	except Exception as e:
		raise HTTPException(status_code=500, detail=f"Report generation failed: {e}")


@app.post("/stt")
async def speech_to_text(file: UploadFile = File(...)):
	"""Accepts an audio file and returns a transcript using SarvamAI."""
	try:
		# Read API key from environment
		api_key = os.getenv("SARVAMAI_API_KEY")
		if not api_key:
			raise HTTPException(status_code=500, detail="SARVAMAI_API_KEY not configured in environment")
		
		# Create SarvamAI client
		client = SarvamAI(api_subscription_key=api_key)
		
		# Call speech_to_text.translate with the audio file
		response = client.speech_to_text.translate(
			file=file.file,
			model="saaras:v2.5"
		)
		
		# Extract transcript from response
		transcript = response.get("transcript") if isinstance(response, dict) else str(response)
		
		return {"transcript": transcript}
	except Exception as e:
		raise HTTPException(status_code=500, detail=f"Speech-to-text failed: {str(e)}")


if __name__ == "__main__":
	uvicorn.run(app, host="0.0.0.0", port=8000)


